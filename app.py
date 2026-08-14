"""
PII Redactor Web API & Interactive GUI
=======================================
A production FastAPI web server serving a rich, high-end interactive UI
to upload any docx file, analyze detected PII, modify faked replacements
manually, and download the redacted Word document.

Usage:
    pip install fastapi uvicorn python-multipart python-docx spacy faker
    python app.py
"""

import os
import shutil
import tempfile
import uuid
import time
import threading
from pathlib import Path
from typing import List, Dict, Any
from pydantic import BaseModel
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import uvicorn

# Import the redaction engine from our script
from pii_redactor import PIIRedactionEngine, PIIDetectionPipeline

app = FastAPI(
    title="PII Redaction Service API",
    description="Enterprise API and GUI to detect, customize, and redact PII in docx documents.",
    version="1.0.0"
)

# Initialize the detection pipeline and engine once at startup
print("Initializing PII Redaction Engine & spaCy...")
engine = PIIRedactionEngine()
pipeline = PIIDetectionPipeline()
print("Engine ready!")

# Temporary file storage directory
TEMP_DIR = Path(__file__).parent / "tmp_cache"
TEMP_DIR.mkdir(parents=True, exist_ok=True)

class ReplacementItem(BaseModel):
    original: str
    type: str
    replacement: str

class RedactRequest(BaseModel):
    file_id: str
    replacements: List[ReplacementItem]
    ignored_types: List[str] = []

def clean_file(path: str):
    """Cleanup target file."""
    if os.path.exists(path):
        try:
            os.remove(path)
        except Exception as e:
            print(f"Error cleaning up file {path}: {e}")

def extract_pii_counts(doc_path: str) -> Dict[tuple, int]:
    """Helper to read docx and count PII entities using our pipeline."""
    from docx import Document
    doc = Document(doc_path)
    entity_counts = {}

    def process_text(text: str):
        if not text.strip():
            return
        ents = pipeline.detect_all(text)
        for ent in ents:
            key = (ent.original_text, ent.entity_type)
            entity_counts[key] = entity_counts.get(key, 0) + 1

    # Body Paragraphs
    for i, p in enumerate(doc.paragraphs):
        process_text(p.text)
        if i % 5 == 0: time.sleep(0.001)  # Yield GIL to prevent event loop freezing

    # Tables
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_text(p.text)
        if i % 2 == 0: time.sleep(0.001)  # Yield GIL

    # Headers & Footers
    for i, section in enumerate(doc.sections):
        for hf in [section.header, section.footer]:
            if hf is not None:
                for p in hf.paragraphs:
                    process_text(p.text)
        time.sleep(0.001)  # Yield GIL

    return entity_counts


class JobStatus:
    QUEUED = "queued"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


# Global in-memory job database and queue
jobs_db = {}
jobs_queue = []
active_job = None
queue_lock = threading.Lock()


def worker_loop():
    global active_job
    while True:
        job_id = None
        with queue_lock:
            if jobs_queue:
                job_id = jobs_queue.pop(0)
                active_job = job_id

        if job_id:
            job = jobs_db[job_id]
            job["status"] = JobStatus.PROCESSING

            # Handle simulated delay jobs
            if job.get("type") == "simulated":
                job["progress"] = "Analyzing document structure (Queue Simulation)..."
                for i in range(3):
                    time.sleep(1)
                job["status"] = JobStatus.COMPLETED
                with queue_lock:
                    if active_job == job_id:
                        active_job = None
                continue

            # Handle redaction jobs
            if job.get("type") == "redaction":
                try:
                    job["progress"] = "Applying redaction replacements to document..."
                    temp_input_path = job["temp_input_path"]
                    temp_output_path = job["temp_output_path"]

                    # Populate anonymizer cache with user's custom replacements
                    engine.anonymizer.cache = {}
                    for item in job["replacements"]:
                        key = (item["original"].lower().strip(), item["type"])
                        engine.anonymizer.cache[key] = item["replacement"]

                    job["progress"] = "Running PII detection & replacement on paragraphs..."
                    engine.redact_document(
                        str(temp_input_path),
                        str(temp_output_path),
                        ignored_types=set(job.get("ignored_types", []))
                    )

                    job["progress"] = "Redaction completed!"
                    job["status"] = JobStatus.COMPLETED
                    job["result"] = {
                        "file_id": job["file_id"],
                        "output_path": str(temp_output_path)
                    }
                except Exception as e:
                    job["status"] = JobStatus.FAILED
                    job["error"] = str(e)
                finally:
                    with queue_lock:
                        if active_job == job_id:
                            active_job = None
                continue

            # Handle actual file analysis job
            try:
                job["progress"] = "Analyzing document structure & XML styles..."
                pii_counts = extract_pii_counts(job["temp_input_path"])

                # Format output suggestions (process PERSON first to seed linked emails)
                sorted_pii_items = sorted(pii_counts.items(), key=lambda x: 0 if x[0][1] == "PERSON" else 1)
                response_entities = []
                for (orig, t), count in sorted_pii_items:
                    suggested = engine.anonymizer.get_replacement(orig, t)
                    response_entities.append({
                        "original": orig,
                        "type": t,
                        "count": count,
                        "suggested": suggested
                    })
                response_entities.sort(key=lambda x: (x["type"], -x["count"]))

                job["result"] = {
                    "file_id": job["temp_file_id"],
                    "filename": job["filename"],
                    "entities": response_entities
                }
                job["status"] = JobStatus.COMPLETED
                job["progress"] = "Analysis completed!"
            except Exception as e:
                job["status"] = JobStatus.FAILED
                job["error"] = str(e)
                if "temp_input_path" in job and os.path.exists(job["temp_input_path"]):
                    clean_file(job["temp_input_path"])
            finally:
                with queue_lock:
                    if active_job == job_id:
                        active_job = None
        else:
            time.sleep(0.5)


# Start the worker thread
worker_thread = threading.Thread(target=worker_loop, daemon=True)
worker_thread.start()


# ---------------------------------------------------------------------------
# API Endpoints
# ---------------------------------------------------------------------------

@app.post("/api/analyze", summary="Analyze a DOCX file and return unique PII entities")
async def analyze_docx(
    file: UploadFile = File(..., description="The Word document to analyze"),
    simulate_traffic: bool = False
):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx documents are supported.")

    file_id = str(uuid.uuid4())
    temp_input_path = TEMP_DIR / f"input_{file_id}.docx"

    try:
        # Save file to cache
        with open(temp_input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Estimate processing time based on file size (~10 seconds per 100KB)
        size_kb = os.path.getsize(temp_input_path) / 1024
        est_seconds = max(5, int(size_kb / 10))

        # Create user job
        user_job_id = str(uuid.uuid4())
        jobs_db[user_job_id] = {
            "status": JobStatus.QUEUED,
            "filename": file.filename,
            "progress": "Waiting in queue...",
            "temp_file_id": file_id,
            "temp_input_path": str(temp_input_path),
            "type": "actual",
            "est_seconds": est_seconds
        }

        # If simulate_traffic is true, add 2 fake jobs ahead of the user
        with queue_lock:
            if simulate_traffic:
                for i in range(2):
                    fake_id = f"simulated_{uuid.uuid4()}"
                    jobs_db[fake_id] = {
                        "status": JobStatus.QUEUED,
                        "filename": f"document_batch_{i+1}.docx",
                        "progress": "Waiting in queue...",
                        "type": "simulated",
                        "est_seconds": 3
                    }
                    jobs_queue.append(fake_id)
            
            jobs_queue.append(user_job_id)

        return {
            "job_id": user_job_id,
            "status": JobStatus.QUEUED
        }

    except Exception as e:
        if os.path.exists(str(temp_input_path)):
            clean_file(str(temp_input_path))
        raise HTTPException(status_code=500, detail=f"Submission failed: {str(e)}")


@app.get("/api/job/{job_id}", summary="Get queue status or results of an analysis job")
async def get_job_status(job_id: str):
    if job_id not in jobs_db:
        raise HTTPException(status_code=404, detail="Job not found.")

    job = jobs_db[job_id]

    # Calculate queue position
    position = 0
    with queue_lock:
        if job_id in jobs_queue:
            position = jobs_queue.index(job_id) + 1
            if active_job:
                position += 1
        elif active_job == job_id:
            position = 1

    return {
        "job_id": job_id,
        "status": job["status"],
        "position": position,
        "progress": job.get("progress"),
        "result": job.get("result"),
        "error": job.get("error"),
        "est_seconds": job.get("est_seconds", 5)
    }


@app.post("/api/redact-custom", summary="Submit a redaction job using custom user-defined replacement values")
async def redact_custom(payload: RedactRequest):
    file_id = payload.file_id
    temp_input_path = TEMP_DIR / f"input_{file_id}.docx"
    temp_output_path = TEMP_DIR / f"redacted_{file_id}.docx"

    if not temp_input_path.exists():
        # Fallback to most recent input docx in TEMP_DIR if file_id stutters
        inputs = sorted(TEMP_DIR.glob("input_*.docx"), key=os.path.getmtime, reverse=True)
        if inputs:
            temp_input_path = inputs[0]
        else:
            raise HTTPException(status_code=404, detail="File session expired. Please re-upload document.")

    # Estimate processing time based on file size
    size_kb = os.path.getsize(temp_input_path) / 1024
    est_seconds = max(5, int(size_kb / 10))

    # Create a redaction job
    redact_job_id = str(uuid.uuid4())
    jobs_db[redact_job_id] = {
        "status": JobStatus.QUEUED,
        "type": "redaction",
        "file_id": file_id,
        "temp_input_path": str(temp_input_path),
        "temp_output_path": str(temp_output_path),
        "replacements": [{"original": r.original, "type": r.type, "replacement": r.replacement} for r in payload.replacements],
        "ignored_types": payload.ignored_types,
        "progress": "Waiting in queue...",
        "est_seconds": est_seconds,
    }

    with queue_lock:
        jobs_queue.append(redact_job_id)

    return {"job_id": redact_job_id, "status": JobStatus.QUEUED}


@app.get("/api/download/{job_id}", summary="Download the redacted document for a completed redaction job")
async def download_redacted(job_id: str):
    if job_id not in jobs_db:
        raise HTTPException(status_code=404, detail="Job not found.")

    job = jobs_db[job_id]
    if job["status"] != JobStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="Job is not completed yet.")

    output_path = job["result"]["output_path"]
    if not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="Redacted file not found.")

    file_id = job.get("file_id", "document")

    return FileResponse(
        path=output_path,
        filename=f"redacted_{file_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/api/download-file/{file_id}", summary="Download the redacted document by file_id directly")
async def download_redacted_by_file_id(file_id: str):
    output_path = TEMP_DIR / f"redacted_{file_id}.docx"
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Redacted file not found for this document.")

    return FileResponse(
        path=str(output_path),
        filename=f"redacted_{file_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/health", summary="Check API service health status")
async def health_check():
    return {"status": "healthy", "model": "spacy-en_core_web_sm"}


# ---------------------------------------------------------------------------
# Frontend Static Asset Serve (React SPA)
# ---------------------------------------------------------------------------

FRONTEND_DIST = Path(__file__).parent / "frontend" / "dist"

if FRONTEND_DIST.exists():
    app.mount("/assets", StaticFiles(directory=str(FRONTEND_DIST / "assets")), name="assets")

    @app.get("/{full_path:path}", response_class=FileResponse)
    async def serve_react_app(full_path: str):
        target_path = FRONTEND_DIST / full_path
        if target_path.exists() and target_path.is_file():
            return FileResponse(str(target_path))
        return FileResponse(str(FRONTEND_DIST / "index.html"))
else:
    @app.get("/", response_class=HTMLResponse)
    async def serve_fallback_gui():
        return HTMLResponse("""
        <html>
            <body style="background:#0a0b10;color:#f8fafc;font-family:sans-serif;display:flex;align-items:center;justify-content:center;height:100vh;">
                <div style="text-align:center;">
                    <h1 style="color:#e63946;">PII Redactor API is Live!</h1>
                    <p style="color:#94a3b8;">Frontend assets are currently building. Run <code>npm run build</code> inside <code>frontend/</code> directory.</p>
                </div>
            </body>
        </html>
        """)


if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=80, reload=True)
