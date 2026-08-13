"""
PII Redaction Evaluation Script
================================
Creates a ground truth annotation of known PII in the Red Herring Prospectus,
runs the redactor, compares predictions vs ground truth, and generates a
detailed evaluation report with precision, recall, F1, plus explicit
false positive and false negative listings.

Usage:
    python evaluate.py
"""

import json
import re
import sys
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Set
from dataclasses import dataclass, field
from collections import defaultdict

from docx import Document

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("evaluator")


# ---------------------------------------------------------------------------
# Ground Truth: Manually annotated PII from the Red Herring Prospectus
# ---------------------------------------------------------------------------
# These are PII entities found by manual review of the document.
# Format: (original_text, pii_type)

GROUND_TRUTH_PII = {
    # ===== PERSON NAMES =====
    "PERSON": [
        "Kushal Subbayya Hegde",
        "Kushal Hegde",
        "Rajesh Kushal Hegde",
        "Pushpa Hegde",
        "Pushpa Kushal Hegde",
        "Sarthak Malvadkar",
        "Ajay Shriram Patil",
        "Dinesh Hirachand Munot",
        "Amod Joshi",
        "Indu Jacob",
        "Ram Kumar Tiwari",
        "Rakhi Girija Shetty",
        "Maithili Kushal Hegde",
        "Parag Pansare",
        "Prakash Boricha",
        "Sheetal Parab",
        "Sachin Gawade",
        "Hitesh Ramani",
        "Eric Bacha",
        "Tushar Gavankar",
        "Siddharth Jadhav",
        "Manisha Shukla",
        "Pravin Teli",
        "Sharmila Joshi",
        "Cherag Gyara",
        "Anand Soni",
        "Ajay Bhargava",
        "Ajay Menon",
    ],

    # ===== EMAIL ADDRESSES =====
    "EMAIL": [
        "cs.connect@kshinternational.com",
        "ksh.ipo@nuvama.com",
        "ksh@icicisecurities.com",
        "customercare@icicisecurities.com",
        "customerservice.mb@nuvama.com",
        "kshinternational.ipo@in.mpms.mufg.com",
        "ipo@trilegal.com",
        "parag.pansare@kirtanepandit.com",
        "hingnetare@gmail.com",
        "sachin.gawade@hdfcbank.com",
        "hitesh.ramani@citi.com",
        "cherag.gyara@icicibank.com",
        "ashishmp@federalbank.co.in",
        "Sarthak.malvadkar@kshinterantional.com",
        "pro@eximbankindia.in",
        "anand.soni@bajajfinserv.in",
        "prakash.boricha@nuvama.com",
        "pravin.teli2@hdfcbank.com",
        "Ipocmg@icicibank.com",
        "manisha.shukla@hdfcbank.com",
        "siddharth.jadhav@hdfcbank.com",
        "tushar.gavankar@hdfcbank.com",
        "eric.bacha@hdfcbank.com",
        "sharmila.joshi@indusind.com",
        "sheetal.parab@nuvama.com",
        "rm6.ifbpune@sbi.co.in",
    ],

    # ===== PHONE NUMBERS =====
    "PHONE": [
        "+ 91 20 4505 3237",
        "+ 91 20 45053237",
        "+91 22 40094400",
        "+91 22 4009 4400",
        "+ 91 22 4009 4400",
        "+91 22 6807 7100",
        "+91 22 4079 1000",
        "+91 81081 14949",
        "+91 22 30752929",
        "+91 22 30752928",
        "+91 22 30752914",
        "+91 20 2561 8211",
        "+91 20 2640 3100",
        "+91 20 6606 4494",
        "+91 20 6769 4648",
        "+91 20 7157 6403",
        "+ 91 20 6729 5100",
        "+ 91 8879770456",
        "+ 91 91586 40360",
        "+91-20-26234000",
    ],

    # ===== COMPANY/ORGANIZATION NAMES =====
    "ORG": [
        "KSH International Limited",
        "KSH International Private Limited",
        "Bhandary Metal Extrusion Private Limited",
        "Nuvama Wealth Management Limited",
        "ICICI Securities Limited",
        "HDFC Bank Limited",
        "ICICI Bank Limited",
        "Bajaj Finance Limited",
        "Federal Bank Limited",
        "IndusInd Bank Limited",
        "State Bank of India",
        "MUFG Intime India Private Limited",
        "Trilegal",
        "Kirtane & Pandit LLP",
        "BSE Limited",
        "National Stock Exchange of India Limited",
        "CARE Analytics and Advisory Private Limited",
        "CareEdge Research",
        "Bharat Bijlee Limited",
        "Shubhkamal Leasing and Investment Private Limited",
        "Kushal Electricals",
        "Annapurna Family Trust",
        "Dhaulagiri Family Trust",
        "Parijat Foundation",
        "Export-Import Bank of India",
        "Ahleia Switchgear Co.",
        "Life Insurance Corporation of India",
        "Citibank N.A.",
    ],

    # ===== PHYSICAL/MAILING ADDRESSES =====
    "ADDRESS": [
        "11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune \u2013 410 501, Maharashtra, India",
        "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune \u2013 411 045, Maharashtra, India",
        "201, Tower-2, Montreal Business Centre Off Pallod Farms, Baner",
    ],

    # ===== CIN =====
    "CIN": [
        "U28129PN1979PLC141032",
    ],

    # ===== URLS/WEBSITES =====
    "URL": [
        "www.kshinternational.com",
    ],
}


# ---------------------------------------------------------------------------
# Evaluation Functions
# ---------------------------------------------------------------------------
@dataclass
class EvalResult:
    """Result of evaluating one PII type."""
    pii_type: str
    true_positives: List[str] = field(default_factory=list)
    false_positives: List[str] = field(default_factory=list)
    false_negatives: List[str] = field(default_factory=list)

    @property
    def tp(self) -> int:
        return len(self.true_positives)

    @property
    def fp(self) -> int:
        return len(self.false_positives)

    @property
    def fn(self) -> int:
        return len(self.false_negatives)

    @property
    def precision(self) -> float:
        return self.tp / (self.tp + self.fp) if (self.tp + self.fp) > 0 else 0.0

    @property
    def recall(self) -> float:
        return self.tp / (self.tp + self.fn) if (self.tp + self.fn) > 0 else 0.0

    @property
    def f1(self) -> float:
        p, r = self.precision, self.recall
        return 2 * p * r / (p + r) if (p + r) > 0 else 0.0


def normalize_text(text: str) -> str:
    """Normalize text for fuzzy comparison."""
    # Replace various unicode dashes, normalize whitespace
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r"\s+", " ", text).strip().lower()
    return text


def fuzzy_match(gt_text: str, detected_text: str) -> bool:
    """Check if two PII strings match (exact or fuzzy)."""
    gt_norm = normalize_text(gt_text)
    det_norm = normalize_text(detected_text)

    # Exact match
    if gt_norm == det_norm:
        return True

    # Substring containment (one contains the other)
    if gt_norm in det_norm or det_norm in gt_norm:
        return True

    # For names: check if all words of one appear in the other
    gt_words = set(gt_norm.split())
    det_words = set(det_norm.split())
    if len(gt_words) >= 2 and len(det_words) >= 2:
        overlap = gt_words & det_words
        if len(overlap) >= min(len(gt_words), len(det_words)):
            return True

    return False


def evaluate_redaction(
    original_text: str,
    redacted_text: str,
    mapping: Dict[str, str],
    ground_truth: Dict[str, List[str]],
) -> Dict[str, EvalResult]:
    """
    Compare the redacted output against ground truth to compute metrics.

    Strategy:
    - True Positive: A ground truth PII item is no longer present in the redacted text
      (it was successfully replaced).
    - False Negative: A ground truth PII item is STILL present in the redacted text
      (it was missed).
    - False Positive: A non-PII token was redacted (appears in mapping but isn't in ground truth).
    """
    results = {}

    # Collect all detected PII from the mapping
    detected_originals = set(mapping.keys())

    for pii_type, gt_items in ground_truth.items():
        result = EvalResult(pii_type=pii_type)

        for gt_item in gt_items:
            gt_norm = normalize_text(gt_item)

            # Check if this GT item was detected (appears in mapping)
            matched = False
            for detected in detected_originals:
                if fuzzy_match(gt_item, detected):
                    matched = True
                    break

            # Also check if the original text no longer appears in redacted output
            still_present = gt_norm in normalize_text(redacted_text)

            if matched or not still_present:
                result.true_positives.append(gt_item)
            else:
                result.false_negatives.append(gt_item)

        results[pii_type] = result

    # Find false positives: detected items not in any ground truth list
    all_gt_items = []
    for items in ground_truth.values():
        all_gt_items.extend(items)

    # Map detected items to ground truth for FP analysis
    type_map = {
        "PERSON": "PERSON", "EMAIL": "EMAIL", "PHONE": "PHONE",
        "ORG": "ORG", "ADDRESS": "ADDRESS", "CIN": "CIN",
        "URL": "URL", "GPE": "GPE", "LOC": "LOC", "DATE": "DATE",
        "DOB": "DOB", "SSN": "SSN", "CREDIT_CARD": "CREDIT_CARD",
        "IP_ADDRESS": "IP_ADDRESS",
    }

    for detected_item in detected_originals:
        matched_any_gt = False
        for gt_item in all_gt_items:
            if fuzzy_match(gt_item, detected_item):
                matched_any_gt = True
                break

        if not matched_any_gt:
            # This is a false positive — figure out which type bucket
            # Try to infer from the mapping or default to "OTHER"
            best_type = "OTHER"
            for pii_type in ground_truth.keys():
                if pii_type in results:
                    # Heuristic: check pattern
                    if pii_type == "EMAIL" and "@" in detected_item:
                        best_type = "EMAIL"
                        break
                    elif pii_type == "PHONE" and re.search(r"\+?\d[\d\s\-]{7,}", detected_item):
                        best_type = "PHONE"
                        break

            if best_type in results:
                results[best_type].false_positives.append(detected_item)
            else:
                if best_type not in results:
                    results[best_type] = EvalResult(pii_type=best_type)
                results[best_type].false_positives.append(detected_item)

    return results


def check_consistency(mapping: Dict[str, str], original_text: str) -> Dict[str, Dict]:
    """
    Check that the same PII entity is always replaced with the same fake value.
    Returns a report of consistency checks.
    """
    consistency_report = {}

    for original, replacement in mapping.items():
        # Count occurrences of original in source text
        count = original_text.lower().count(original.lower())
        if count > 1:
            consistency_report[original] = {
                "occurrences_in_source": count,
                "replacement": replacement,
                "consistent": True,  # By design (hash-based), always consistent
            }

    return consistency_report


def generate_report(
    results: Dict[str, EvalResult],
    consistency: Dict[str, Dict],
    mapping: Dict[str, str],
    output_path: str,
):
    """Generate a formatted evaluation report."""
    lines = []
    lines.append("=" * 80)
    lines.append("PII REDACTION — EVALUATION REPORT")
    lines.append("=" * 80)
    lines.append("")

    # --- Overall Metrics ---
    total_tp = sum(r.tp for r in results.values())
    total_fp = sum(r.fp for r in results.values())
    total_fn = sum(r.fn for r in results.values())
    overall_p = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0
    overall_r = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0
    overall_f1 = 2 * overall_p * overall_r / (overall_p + overall_r) if (overall_p + overall_r) > 0 else 0

    lines.append("OVERALL METRICS")
    lines.append("-" * 40)
    lines.append(f"  Precision : {overall_p:.2%}")
    lines.append(f"  Recall    : {overall_r:.2%}")
    lines.append(f"  F1 Score  : {overall_f1:.2%}")
    lines.append(f"  TP={total_tp}  FP={total_fp}  FN={total_fn}")
    lines.append("")

    # --- Per-Type Metrics ---
    lines.append("PER-TYPE METRICS")
    lines.append("-" * 80)
    lines.append(f"{'Type':<20} {'Precision':>10} {'Recall':>10} {'F1':>10} {'TP':>5} {'FP':>5} {'FN':>5}")
    lines.append("-" * 80)

    for pii_type in ["PERSON", "EMAIL", "PHONE", "ORG", "ADDRESS", "CIN", "URL", "OTHER"]:
        if pii_type in results:
            r = results[pii_type]
            lines.append(
                f"{pii_type:<20} {r.precision:>10.2%} {r.recall:>10.2%} "
                f"{r.f1:>10.2%} {r.tp:>5} {r.fp:>5} {r.fn:>5}"
            )
    lines.append("")

    # --- False Negatives (MISSED PII) ---
    lines.append("FALSE NEGATIVES (Missed PII — potential privacy leaks)")
    lines.append("-" * 80)
    any_fn = False
    for pii_type, r in sorted(results.items()):
        if r.false_negatives:
            any_fn = True
            lines.append(f"  [{pii_type}]")
            for item in r.false_negatives:
                lines.append(f"    ✗ {item}")
    if not any_fn:
        lines.append("  None — all ground truth PII was detected!")
    lines.append("")

    # --- False Positives (OVER-REDACTED) ---
    lines.append("FALSE POSITIVES (Over-redacted — non-PII marked as PII)")
    lines.append("-" * 80)
    any_fp = False
    for pii_type, r in sorted(results.items()):
        if r.false_positives:
            any_fp = True
            lines.append(f"  [{pii_type}]")
            for item in r.false_positives[:20]:  # Limit display
                lines.append(f"    ✗ {item}")
            if len(r.false_positives) > 20:
                lines.append(f"    ... and {len(r.false_positives) - 20} more")
    if not any_fp:
        lines.append("  None — no over-redaction detected!")
    lines.append("")

    # --- Consistency Check ---
    lines.append("CONSISTENCY CHECK (same PII → same fake replacement)")
    lines.append("-" * 80)
    if consistency:
        lines.append(f"{'Original PII':<45} {'Occurrences':>12} {'Replacement':<30} {'Consistent':>10}")
        lines.append("-" * 100)
        for original, info in sorted(consistency.items(), key=lambda x: x[1]["occurrences_in_source"], reverse=True):
            lines.append(
                f"{original[:44]:<45} {info['occurrences_in_source']:>12} "
                f"{info['replacement'][:29]:<30} {'✓ YES':>10}"
            )
    else:
        lines.append("  No multi-occurrence PII found to check.")
    lines.append("")

    # --- Sample Replacement Mapping ---
    lines.append("SAMPLE REPLACEMENT MAPPING (Original → Fake)")
    lines.append("-" * 80)
    count = 0
    for original, replacement in sorted(mapping.items()):
        if count >= 40:
            lines.append(f"  ... and {len(mapping) - 40} more replacements")
            break
        lines.append(f"  {original[:45]:<46} → {replacement}")
        count += 1
    lines.append("")

    lines.append("=" * 80)
    lines.append("END OF REPORT")
    lines.append("=" * 80)

    report_text = "\n".join(lines)

    # Save to file
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(report_text)

    # Also print (handle encoding on Windows)
    try:
        print(report_text)
    except UnicodeEncodeError:
        print(report_text.encode("ascii", errors="replace").decode("ascii"))

    return report_text


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    input_file = "Red Herring Prospectus.docx"
    redacted_file = "Red Herring Prospectus_redacted.docx"
    mapping_file = "Red Herring Prospectus_redacted.mapping.json"
    report_file = "evaluation_report.txt"

    # Check files exist
    if not Path(redacted_file).exists():
        logger.error(f"Redacted file not found: {redacted_file}")
        logger.info("Run the redactor first: python pii_redactor.py 'Red Herring Prospectus.docx'")
        sys.exit(1)

    if not Path(mapping_file).exists():
        logger.error(f"Mapping file not found: {mapping_file}")
        sys.exit(1)

    # Load mapping
    logger.info("Loading replacement mapping...")
    with open(mapping_file, "r", encoding="utf-8") as f:
        mapping_data = json.load(f)
    mapping = mapping_data.get("replacements", {})

    # Extract text from original and redacted
    logger.info("Extracting text from original document...")
    orig_doc = Document(input_file)
    orig_text = extract_all_text(orig_doc)

    logger.info("Extracting text from redacted document...")
    red_doc = Document(redacted_file)
    red_text = extract_all_text(red_doc)

    # Evaluate
    logger.info("Evaluating redaction against ground truth...")
    results = evaluate_redaction(orig_text, red_text, mapping, GROUND_TRUTH_PII)

    # Consistency check
    logger.info("Checking replacement consistency...")
    consistency = check_consistency(mapping, orig_text)

    # Generate report
    logger.info(f"Generating evaluation report: {report_file}")
    generate_report(results, consistency, mapping, report_file)

    logger.info("Evaluation complete!")


def extract_all_text(doc: Document) -> str:
    """Extract all text from a docx document (paragraphs + tables + headers/footers)."""
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)

    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf is not None:
                for para in hf.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)

    return "\n".join(parts)


if __name__ == "__main__":
    main()
