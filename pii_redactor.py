"""
PII Redaction Tool
==================
A hybrid PII detection and redaction engine using spaCy NER + regex patterns
to detect and replace Personally Identifiable Information (PII) in .docx files.

Approach:
  1. spaCy NER (en_core_web_sm) for contextual entity recognition (PERSON, ORG, GPE, LOC, DATE)
  2. Regex patterns for structured PII (emails, phones, SSNs, credit cards, IPs, DOBs)
  3. Contextual heuristics for addresses (Indian PIN codes, multi-line address blocks)
  4. Faker library for consistent, deterministic fake data generation

Author: PII Redaction Engine
"""

import re
import sys
import json
import hashlib
import logging
import argparse
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Set
from dataclasses import dataclass, field, asdict
from copy import deepcopy

import spacy
from faker import Faker
from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Logging Configuration
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("pii_redactor")


def detect_best_spacy_model() -> str:
    """Helper to detect the best installed spaCy model in the environment."""
    for model in ["en_core_web_trf", "en_core_web_lg", "en_core_web_md"]:
        try:
            if spacy.util.is_package(model):
                return model
        except Exception:
            pass
    return "en_core_web_sm"


# ---------------------------------------------------------------------------
# Data Classes
# ---------------------------------------------------------------------------
@dataclass
class PIIEntity:
    """Represents a detected PII entity with its span and metadata."""
    start: int
    end: int
    entity_type: str
    original_text: str
    replacement: str = ""
    source: str = "unknown"        # "regex", "ner", "heuristic"
    confidence: float = 1.0

    def __hash__(self):
        return hash((self.start, self.end, self.entity_type))

    def __eq__(self, other):
        return (self.start, self.end, self.entity_type) == (other.start, other.end, other.entity_type)


@dataclass
class RedactionStats:
    """Tracks statistics about the redaction process."""
    total_entities: int = 0
    by_type: Dict[str, int] = field(default_factory=dict)
    by_source: Dict[str, int] = field(default_factory=dict)
    replacement_map: Dict[str, str] = field(default_factory=dict)

    def record(self, entity: PIIEntity):
        self.total_entities += 1
        self.by_type[entity.entity_type] = self.by_type.get(entity.entity_type, 0) + 1
        self.by_source[entity.source] = self.by_source.get(entity.source, 0) + 1
        if entity.original_text not in self.replacement_map:
            self.replacement_map[entity.original_text] = entity.replacement


# ---------------------------------------------------------------------------
# Consistent Fake Data Generator
# ---------------------------------------------------------------------------
class ConsistentAnonymizer:
    """
    Generates fake PII replacements deterministically.
    Same input always maps to the same fake output across the entire document,
    preserving entity relationships and document coherence.
    """

    def __init__(self, locale: str = "en_US", salt: str = "pii_redaction_salt_2026"):
        self.fake = Faker(locale)
        self.salt = salt
        self.cache: Dict[Tuple[str, str], str] = {}

    def _seed_for(self, text: str, entity_type: str) -> int:
        """Deterministic seed from (text, type) pair."""
        key = f"{text.lower().strip()}|{entity_type}|{self.salt}"
        return int(hashlib.md5(key.encode("utf-8")).hexdigest(), 16) % (2**32)

    def clear_cache(self):
        """Clear the mapping cache to prevent memory leaks and cross-document state bleed."""
        self.cache.clear()

    def get_replacement(self, original: str, entity_type: str) -> str:
        """Return a consistent fake replacement for the given PII entity."""
        cache_key = (original.lower().strip(), entity_type)
        if cache_key in self.cache:
            return self.cache[cache_key]

        seed = self._seed_for(original, entity_type)
        self.fake.seed_instance(seed)

        replacement = self._generate(original, entity_type)
        self.cache[cache_key] = replacement
        return replacement

    def _generate(self, original: str, entity_type: str) -> str:
        """Generate a type-appropriate fake value."""
        generators = {
            "PERSON": lambda: self.fake.name(),
            "EMAIL": lambda: self._generate_linked_email(original),
            "PHONE": lambda: self._fake_phone(original),
            "ORG": lambda: self.fake.company(),
            "ADDRESS": lambda: self.fake.address().replace("\n", ", "),
            "SSN": lambda: self.fake.ssn(),
            "CREDIT_CARD": lambda: self.fake.credit_card_number(),
            "DOB": lambda: str(self.fake.date_of_birth(minimum_age=18, maximum_age=70)),
            "DATE": lambda: str(self.fake.date_between(start_date="-10y", end_date="today")),
            "IP_ADDRESS": lambda: self.fake.ipv4(),
            "GPE": lambda: self.fake.city(),
            "LOC": lambda: self.fake.city(),
            "URL": lambda: f"www.{self.fake.domain_name()}",
            "CIN": lambda: self._fake_cin(),
            "DIN": lambda: str(self.fake.random_int(min=10000000, max=99999999)),
            "PIN_CODE": lambda: str(self.fake.random_int(min=100000, max=999999)),
        }
        gen = generators.get(entity_type, lambda: f"[{entity_type}_REDACTED]")
        return gen()

    def _generate_linked_email(self, original: str) -> str:
        """Attempt to generate a fake email linked to a person name in the cache."""
        email_lower = original.lower().strip()
        if "@" not in email_lower:
            return self.fake.email()

        username, domain = email_lower.split("@", 1)
        # Clean username into alphanumeric parts
        username_parts = [w for w in re.split(r"[^a-z0-9]", username) if len(w) > 2]

        # Check cache for any PERSON replacement that matches username parts
        best_match = None
        for (cached_orig, cached_type), replacement in self.cache.items():
            if cached_type == "PERSON":
                # Split original name into parts
                name_parts = [w for w in re.split(r"[^a-z0-9]", cached_orig.lower()) if len(w) > 2]
                overlap = False
                for up in username_parts:
                    for np in name_parts:
                        # Direct match or substring match (e.g. rashhi and rashi)
                        if up == np or up in np or np in up:
                            overlap = True
                            break
                    if overlap:
                        break
                if overlap:
                    best_match = replacement
                    break

        if best_match:
            # Clean replacement name into an email username (e.g. "John Doe" -> "john.doe")
            clean_name = re.sub(r"[^a-z0-9\s]", "", best_match.lower())
            email_username = ".".join(clean_name.split())
            generic_domains = {"gmail.com", "yahoo.com", "outlook.com", "hotmail.com", "icloud.com"}
            use_domain = domain if domain in generic_domains else "example.com"
            return f"{email_username}@{use_domain}"

        return self.fake.email()

    def _fake_phone(self, original: str) -> str:
        """Generate a fake phone preserving the original format."""
        digits = re.sub(r"\D", "", original)
        if original.startswith("+"):
            country = "+91 "
            local_digits = "".join([str(self.fake.random_digit()) for _ in range(10)])
            # Try to preserve formatting
            if " " in original[3:]:
                parts = original[3:].strip().split()
                fake_parts = []
                pos = 0
                for part in parts:
                    clean = re.sub(r"\D", "", part)
                    fake_parts.append(local_digits[pos:pos + len(clean)])
                    pos += len(clean)
                return country + " ".join(fake_parts)
            return country + local_digits
        else:
            return "".join([str(self.fake.random_digit()) for _ in range(len(digits))])

    def _fake_cin(self) -> str:
        """Generate a fake Corporate Identity Number."""
        letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        return (
            self.fake.random_element(letters)
            + str(self.fake.random_int(min=10000, max=99999))
            + self.fake.random_element(["MH", "DL", "KA", "TN", "GJ", "RJ"])
            + str(self.fake.random_int(min=1960, max=2024))
            + self.fake.random_element(["PLC", "PTC", "GAP"])
            + str(self.fake.random_int(min=100000, max=999999))
        )


# ---------------------------------------------------------------------------
# Regex-Based PII Detectors
# ---------------------------------------------------------------------------
class RegexDetector:
    """
    Detects structured PII using regex patterns.
    Each pattern is validated with optional post-processing to reduce false positives.
    """

    # Common false-positive terms that should NOT be redacted as organizations
    ORG_BLOCKLIST = {
        "companies act", "income tax act", "securities", "sebi", "rbi",
        "government", "parliament", "lok sabha", "rajya sabha",
    }

    def __init__(self):
        self.patterns = self._build_patterns()

    def _build_patterns(self) -> List[Dict]:
        """Define all regex patterns for structured PII detection."""
        return [
            # Email Addresses
            {
                "type": "EMAIL",
                "pattern": re.compile(
                    r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,7}\b"
                ),
                "validator": None,
                "confidence": 0.99,
            },
            # Indian Phone Numbers (+91, various formats)
            {
                "type": "PHONE",
                "pattern": re.compile(
                    r"(?:\+\s*91[\-\s]*|091[\-\s]*)"    # +91 prefix or 091 prefix
                    r"\d[\d\-\s]{8,13}\d"               # 10 digits with optional separators
                ),
                "validator": self._validate_phone,
                "confidence": 0.90,
            },
            # Landline with STD code
            {
                "type": "PHONE",
                "pattern": re.compile(
                    r"(?:\+\s*91[\-\s]*)?"
                    r"(?:0\d{2,4}[\-\s])"               # STD code
                    r"\d[\d\-\s]{6,9}\d"                 # local number
                ),
                "validator": self._validate_phone,
                "confidence": 0.85,
            },
            # US Social Security Numbers (XXX-XX-XXXX)
            {
                "type": "SSN",
                "pattern": re.compile(
                    r"\b(?!000|666|9\d{2})\d{3}[\-\s]?"
                    r"(?!00)\d{2}[\-\s]?"
                    r"(?!0000)\d{4}\b"
                ),
                "validator": self._validate_ssn,
                "confidence": 0.80,
            },
            # Credit Card Numbers (Visa, MC, Amex, Discover)
            {
                "type": "CREDIT_CARD",
                "pattern": re.compile(
                    r"\b(?:4\d{3}|5[1-5]\d{2}|3[47]\d{2}|6(?:011|5\d{2}))"
                    r"[\-\s]?\d{4}[\-\s]?\d{4}[\-\s]?\d{3,4}\b"
                ),
                "validator": self._validate_luhn,
                "confidence": 0.95,
            },
            # Date of Birth (DD/MM/YYYY, MM/DD/YYYY, DD-MM-YYYY, etc.)
            {
                "type": "DOB",
                "pattern": re.compile(
                    r"\b(?:0?[1-9]|[12]\d|3[01])"
                    r"[\/\-\.]"
                    r"(?:0?[1-9]|1[0-2])"
                    r"[\/\-\.]"
                    r"(?:19|20)\d{2}\b"
                ),
                "validator": None,
                "confidence": 0.75,
            },
            # IPv4 Addresses
            {
                "type": "IP_ADDRESS",
                "pattern": re.compile(
                    r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}"
                    r"(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b"
                ),
                "validator": None,
                "confidence": 0.95,
            },
            # IPv6 Addresses
            {
                "type": "IP_ADDRESS",
                "pattern": re.compile(
                    r"\b(?:[0-9a-fA-F]{1,4}:){7}[0-9a-fA-F]{1,4}\b"
                ),
                "validator": None,
                "confidence": 0.95,
            },
            # Corporate Identity Number (CIN) - Indian specific
            {
                "type": "CIN",
                "pattern": re.compile(
                    r"\b[A-Z]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b"
                ),
                "validator": None,
                "confidence": 0.95,
            },
            # Website URLs
            {
                "type": "URL",
                "pattern": re.compile(
                    r"(?:https?://)?www\.[a-zA-Z0-9\-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?"
                ),
                "validator": None,
                "confidence": 0.90,
            },
            # Company/Organization Names (standard business suffixes)
            {
                "type": "ORG",
                "pattern": re.compile(
                    r"\b[A-Z][A-Za-z0-9\-]*"
                    r"(?:\s+(?:[A-Z][A-Za-z0-9\-]*|and|of|for|in|the|&|\([A-Za-z0-9\-]+\)))*"
                    r"\s+(?:Private\s+Limited|Limited|Pvt\s+Ltd|Ltd|LLP|Co|Corp|Corporation|Inc|Private|Ltd\.|Co\.)\b"
                ),
                "validator": None,
                "confidence": 0.95,
            },
        ]

    def detect(self, text: str) -> List[PIIEntity]:
        """Detect all regex-based PII entities in the text."""
        entities = []
        for spec in self.patterns:
            for match in spec["pattern"].finditer(text):
                matched_text = match.group().strip()
                if spec["validator"] and not spec["validator"](matched_text):
                    continue
                entities.append(PIIEntity(
                    start=match.start(),
                    end=match.end(),
                    entity_type=spec["type"],
                    original_text=matched_text,
                    source="regex",
                    confidence=spec["confidence"],
                ))
        return entities

    # -- Validators --

    @staticmethod
    def _validate_phone(text: str) -> bool:
        """Validate phone number has correct digit count."""
        digits = re.sub(r"\D", "", text)
        # Remove country code if present
        if digits.startswith("91") and len(digits) > 10:
            digits = digits[2:]
        if digits.startswith("0"):
            digits = digits[1:]
        return 7 <= len(digits) <= 12

    @staticmethod
    def _validate_ssn(text: str) -> bool:
        """Basic SSN validation - avoid matching year ranges, monetary amounts etc."""
        digits = re.sub(r"\D", "", text)
        if len(digits) != 9:
            return False
        # Avoid common false positives
        if digits.startswith("000") or digits.startswith("666"):
            return False
        return True

    @staticmethod
    def _validate_luhn(text: str) -> bool:
        """Luhn algorithm checksum validation for credit cards."""
        digits = [int(c) for c in text if c.isdigit()]
        if len(digits) < 13 or len(digits) > 19:
            return False
        checksum = 0
        for i, d in enumerate(reversed(digits)):
            if i % 2 == 1:
                d *= 2
                if d > 9:
                    d -= 9
            checksum += d
        return checksum % 10 == 0


# ---------------------------------------------------------------------------
# NER-Based PII Detector (spaCy)
# ---------------------------------------------------------------------------
class NERDetector:
    """
    Uses spaCy Named Entity Recognition to detect contextual PII
    like person names, organization names, and geographical locations.
    """

    # Map spaCy labels to our PII types
    # NOTE: DATE, GPE, LOC are excluded — dates are not PII in financial
    # documents, and geographic names (cities/states/countries) are public info.
    LABEL_MAP = {
        "PERSON": "PERSON",
        "ORG": "ORG",
    }

    # Comprehensive blocklist of terms that should NEVER be redacted.
    # This is critical for legal/financial docs where domain terms get
    # mis-classified as PERSON or ORG by spaCy.
    FALSE_POSITIVE_NAMES = {
        # -- Regulatory & Government Bodies --
        "india", "indian", "sebi", "rbi", "bse", "nse", "nsdl", "cdsl",
        "government", "government of india", "republic of india",
        "parliament", "lok sabha", "rajya sabha", "supreme court",
        "high court", "central government", "state government",
        "registrar of companies", "roc", "ministry",
        # -- Legal / Statutory Terms --
        "act", "section", "chapter", "article", "regulation", "rule",
        "schedule", "annexure", "form", "table", "part", "clause",
        "sub-section", "proviso", "ordinance", "notification", "circular",
        "companies act", "income tax act", "sebi act", "fema",
        "scra", "scrr", "depositories act", "indian stamp act",
        "sebi icdr regulations", "ind as", "ifrs", "us gaap", "indian gaap",
        # -- Financial / IPO Terms --
        "equity", "share", "shareholder", "debenture", "bond",
        "offer", "prospectus", "ipo", "red herring", "book",
        "the offer", "fresh issue", "net proceeds", "offer price",
        "bid", "bidder", "bidders", "allotment", "allottee",
        "fiscal", "quarter", "annual", "year", "month", "day",
        "risk factors", "our business", "our company", "the company",
        "board of directors", "board", "directors",
        "anchor investor", "anchor investors",
        "qualified institutional", "non-institutional",
        "retail individual", "eligible employee",
        "promoter", "promoters", "promoter group",
        "mutual funds", "mutual fund",
        "asba", "asba account", "asba bidder", "asba bidders", "asba forms",
        "upi", "upi mechanism", "upi bidders", "upi mandate",
        "neft", "rtgs", "imps",
        # -- Financial Metrics --
        "ebitda", "ebitda margin", "net debt", "roe", "roce",
        "revenue", "profit", "loss", "income", "expenses",
        "non-gaap", "non-gaap measures",
        # -- Document Section Headers --
        "definitions and abbreviations", "forward-looking statements",
        "summary of the offer document", "general information",
        "capital structure", "objects of the offer",
        "basis for the offer price", "statement of special tax benefits",
        "industry overview", "outstanding litigation",
        "management discussion", "financial information",
        "restated financial statements",
        # -- Date/Time Terms --
        "march", "april", "may", "june", "july", "august", "september",
        "october", "november", "december", "january", "february",
        "monday", "tuesday", "wednesday", "thursday", "friday",
        "saturday", "sunday",
        # -- Common role titles (not names) --
        "director", "chairman", "secretary", "officer", "manager",
        "auditor", "partner", "investor", "shareholder",
        "managing director", "whole-time director", "independent director",
        "company secretary", "compliance officer", "chief financial officer",
        "chief executive officer",
        # -- Geographic (public, not personal) --
        "maharashtra", "karnataka", "delhi", "mumbai", "pune",
        "kolkata", "chennai", "bengaluru", "hyderabad", "ahmedabad",
        "united states", "usa", "u.s.", "europe", "eu", "sweden",
        "china", "japan", "germany", "france", "uk",
        "bangalore", "new delhi", "tamil nadu", "gujarat",
        "rajasthan", "bombay", "chakan", "khed", "baner",
        # -- Generic terms spaCy often misclassifies --
        "crore", "crores", "lakh", "lakhs", "rupees", "rs", "inr",
        "per annum", "per cent", "percent",
        "floor", "tower", "block", "unit", "phase",
    }

    def __init__(self, model_name: Optional[str] = None):
        if model_name is None:
            model_name = detect_best_spacy_model()
        logger.info(f"Loading spaCy model: {model_name}")
        try:
            self.nlp = spacy.load(model_name)
        except OSError:
            logger.warning(f"Model {model_name} not found. Downloading...")
            spacy.cli.download(model_name)
            self.nlp = spacy.load(model_name)
        # Increase max length for large documents
        self.nlp.max_length = 500_000

    def detect(self, text: str) -> List[PIIEntity]:
        """Detect NER-based PII entities."""
        entities = []

        # Process in chunks if text is very large
        if len(text) > 100_000:
            chunks = self._chunk_text(text, chunk_size=50_000, overlap=500)
        else:
            chunks = [(0, text)]

        for offset, chunk in chunks:
            doc = self.nlp(chunk)
            for ent in doc.ents:
                label = self.LABEL_MAP.get(ent.label_)
                if not label:
                    continue

                ent_text = ent.text.strip()
                ent_lower = ent_text.lower()

                # Remap misclassified entities to ORG or ADDRESS based on keywords
                words = ent_lower.split()
                cleaned_words = [w.strip(",.-#/") for w in words]

                address_indicators = {
                    "apartment", "apartments", "marg", "path", "road", "street", "lane",
                    "building", "house", "plaza", "villa", "villas", "chambers", "estate",
                    "estates", "heights", "park", "square", "colony", "nagar", "bhavan",
                    "bungalow", "society", "floor", "floors", "wing", "wings", "flat",
                    "flats", "block", "blocks", "sector", "sectors", "cross", "main",
                    "gymkhana"
                }

                if any(w in address_indicators for w in cleaned_words):
                    label = "ADDRESS"
                elif label == "PERSON":
                    org_indicators = {
                        "limited", "private", "ltd", "pvt", "llp", "inc", "corp",
                        "bank", "trust", "fund", "company", "group", "association"
                    }
                    if any(w in org_indicators for w in cleaned_words):
                        label = "ORG"

                # Skip common false positives
                if self._is_false_positive(ent_text, label):
                    continue

                entities.append(PIIEntity(
                    start=offset + ent.start_char,
                    end=offset + ent.end_char,
                    entity_type=label,
                    original_text=ent_text,
                    source="ner",
                    confidence=0.80 if label == "PERSON" else 0.70,
                ))

        return entities

    # Set of common English/business/financial words to filter out non-PII terms
    COMMON_WORDS = {
        "act", "action", "activities", "addition", "address", "addresses", "agreement", "agreements",
        "allotment", "allotments", "amount", "amounts", "analysis", "annexure", "application", "applications",
        "approval", "approvals", "asset", "assets", "association", "audit", "audits", "auditor", "auditors",
        "authority", "authorities", "balance", "balances", "bank", "banks", "banking", "basis", "bid", "bids",
        "bidder", "bidders", "board", "boards", "bond", "bonds", "book", "books", "branch", "branches",
        "business", "businesses", "capital", "case", "cases", "category", "categories", "center", "centers",
        "centre", "centres", "certificate", "certificates", "certification", "certifications", "chapter", "chapters",
        "circular", "circulars", "class", "classes", "clause", "clauses", "code", "codes", "colony", "colonies",
        "commission", "commissions", "committee", "committees", "company", "companies", "compliance", "condition", "conditions",
        "contact", "contacts", "contract", "contracts", "cost", "costs", "credit", "credits", "currency", "currencies",
        "customer", "customers", "date", "dates", "day", "days", "debt", "debts", "decision", "decisions",
        "declaration", "declarations", "defaulter", "defaulters", "definition", "definitions", "demand", "demands",
        "deposit", "deposits", "depository", "depositories", "description", "descriptions", "designation", "designations",
        "development", "developments", "director", "directors", "disclosure", "disclosures", "discrepancy", "discrepancies",
        "district", "districts", "document", "documents", "draft", "drafts", "duty", "duties", "earnings",
        "east", "west", "north", "south", "education", "employee", "employees", "employer", "employers",
        "employment", "engine", "engines", "enterprise", "enterprises", "entity", "entities", "equity", "equities",
        "escrow", "establishment", "establishments", "estate", "estates", "evaluation", "evaluations", "examination", "examinations",
        "exchange", "exchanges", "executive", "executives", "expense", "expenses", "facility", "facilities",
        "factor", "factors", "family", "families", "fee", "fees", "finance", "finances", "financial",
        "firm", "firms", "fiscal", "fiscals", "floor", "floors", "form", "forms", "format", "formats",
        "foundation", "foundations", "fund", "funds", "general", "government", "governments", "group", "groups",
        "guideline", "guidelines", "history", "hour", "hours", "house", "houses", "identification", "implication", "implications",
        "income", "incomes", "incorporation", "index", "indexes", "indices", "industry", "industries",
        "information", "infrastructure", "initial", "institution", "institutions", "institutional", "insurance", "insurances",
        "interest", "interests", "intermediary", "intermediaries", "inventory", "inventories", "investment", "investments",
        "investor", "investors", "issue", "issues", "issuer", "issuers", "item", "items", "joint", "jurisdiction", "jurisdictions",
        "key", "land", "lands", "law", "laws", "lead", "leads", "lease", "leases", "leasehold", "legislation", "legislations",
        "letter", "letters", "liability", "liabilities", "license", "licenses", "limit", "limits", "litigation", "litigations",
        "loan", "loans", "location", "locations", "log", "logs", "loss", "losses", "management", "manager", "managers",
        "mandate", "mandates", "margin", "margins", "market", "markets", "material", "materials", "measure", "measures",
        "meeting", "meetings", "member", "members", "memorandum", "methods", "methodology", "methodologies",
        "metric", "metrics", "million", "millions", "ministry", "ministries", "minute", "minutes", "month", "months",
        "mutual", "name", "names", "net", "news", "newspaper", "newspapers", "nominee", "nominees", "note", "notes",
        "notice", "notices", "notification", "notifications", "number", "numbers", "object", "objects", "obligation", "obligations",
        "offer", "offers", "office", "offices", "officer", "officers", "opening", "openings", "operating", "operation", "operations",
        "operational", "opinion", "opinions", "option", "options", "order", "orders", "ordinance", "ordinances",
        "organization", "organizations", "original", "originals", "outstanding", "owner", "owners", "ownership", "ownerships",
        "page", "pages", "paper", "papers", "part", "parts", "partner", "partners", "partnership", "partnerships",
        "payment", "payments", "pension", "pensions", "percent", "percentage", "percentages", "period", "periods",
        "person", "persons", "personnel", "phase", "phases", "phone", "phones", "place", "places", "plan", "plans",
        "plot", "plots", "policy", "policies", "portion", "portions", "position", "positions", "power", "powers",
        "practice", "practices", "prepayment", "price", "prices", "pricing", "principal", "principals",
        "private", "procedure", "procedures", "proceed", "proceeds", "process", "processes", "processing",
        "product", "products", "profit", "profits", "program", "programs", "programme", "programmes",
        "project", "projects", "promoter", "promoters", "proposal", "proposals", "prospectus", "prospectuses",
        "provision", "provisions", "public", "publication", "publications", "purpose", "purposes", "qualification", "qualifications",
        "qualified", "quarter", "quarters", "rating", "ratings", "ratio", "ratios", "reconciliation", "reconciliations",
        "record", "records", "redaction", "redactions", "reference", "references", "refund", "refunds",
        "registrar", "registrars", "registration", "registrations", "regulation", "regulations", "regulatory", "relation", "relations",
        "relationship", "relationships", "repayment", "repayments", "report", "reports", "reporting", "representative", "representatives",
        "requirement", "requirements", "research", "reserve", "reserves", "resolution", "resolutions", "responsibility", "responsibilities",
        "restated", "restriction", "restrictions", "result", "results", "retail", "return", "returns",
        "revenue", "revenues", "risk", "risks", "road", "roads", "role", "roles", "rule", "rules",
        "running", "salary", "salaries", "sale", "sales", "schedule", "schedules", "scheme", "schemes",
        "school", "schools", "search", "searches", "season", "seasons", "second", "seconds", "secretary", "secretaries",
        "section", "sections", "sector", "sectors", "securities", "security", "selling", "seniors", "series",
        "service", "services", "settlement", "settlements", "share", "shares", "shareholder", "shareholders",
        "shaping", "she", "shift", "shifts", "shoe", "shoes", "signature", "signatures", "signatory", "signatories",
        "size", "sizes", "slip", "slips", "society", "societies", "software", "softwares", "solution", "solutions",
        "source", "sources", "space", "spaces", "special", "specials", "split", "splits", "staff", "staffs",
        "stamp", "stamps", "standard", "standards", "state", "states", "statement", "statements", "station", "stations",
        "status", "statuses", "statute", "statutes", "statutory", "stock", "stocks", "street", "streets",
        "structure", "structures", "study", "studies", "sub-section", "sub-sections", "subject", "subjects",
        "submission", "submissions", "subscriber", "subscribers", "subscription", "subscriptions", "subsidiary", "subsidiaries",
        "summary", "summaries", "supervision", "supervisions", "supplement", "supplements", "surveillance", "survey", "surveys",
        "syndicate", "syndicates", "system", "systems", "table", "tables", "tax", "taxes", "taxation",
        "team", "teams", "technology", "technologies", "telephone", "telephones", "term", "terms",
        "territory", "territories", "test", "tests", "text", "texts", "time", "times", "title", "titles",
        "total", "totals", "tower", "towers", "trade", "trades", "transaction", "transactions", "transfer", "transfers",
        "transformation", "transformations", "transition", "transitions", "treasury", "treasuries", "treaty", "treaties",
        "trust", "trusts", "trustee", "trustees", "type", "types", "underwriter", "underwriters", "underwriting",
        "unit", "units", "university", "universities", "usage", "usages", "validation", "validations",
        "value", "values", "variation", "variations", "vehicle", "vehicles", "venture", "ventures",
        "verification", "verifications", "version", "versions", "view", "views", "volume", "volumes",
        "vote", "votes", "voting", "votings", "way", "ways", "web", "website", "websites",
        "week", "weeks", "weight", "weights", "whole-time", "will", "witness", "witnesses",
        "word", "words", "work", "works", "working", "write", "year", "years", "zone", "zones",
        "eligible", "individual", "registered", "first", "second", "third", "last", "new", "old",
        "top", "bottom", "high", "low", "great", "small", "major", "minor", "national", "international",
        "central", "regional", "local", "public", "private", "corporate", "social", "mutual",
        "the", "and", "for", "with", "from", "into", "over", "under", "above", "below", "between",
        "through", "during", "before", "after", "opposite", "near", "next", "along", "about", "against",
        "other", "another", "some", "any", "each", "every", "all", "both", "neither", "either",
        "our", "your", "their", "his", "her", "its", "my", "green", "shoe", "option", "options"
    }

    def _is_all_common_words(self, text: str) -> bool:
        """Return True if the text contains only common English/financial terms."""
        # Find all word tokens
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        if not words:
            return True # purely numbers/punctuation
        
        # Roman numerals pattern
        roman_numerals = re.compile(r'^[ivxldcm]+$')
        for w in words:
            # If word is not in COMMON_WORDS, not a roman numeral, and not a single letter
            if w not in self.COMMON_WORDS and not roman_numerals.match(w) and len(w) > 1:
                return False
        return True

    def _is_false_positive(self, text: str, label: str) -> bool:
        """Check if detected entity is a known false positive."""
        normalized = text.lower().strip()

        # Direct blocklist match
        if normalized in self.FALSE_POSITIVE_NAMES:
            return True

        # Check if the text consists entirely of common dictionary words
        if self._is_all_common_words(text):
            return True

        # --- Length / structural filters ---

        # Skip very short entities (likely noise)
        if len(normalized) < 3:
            return True

        # Skip entities that are purely numeric or mostly digits
        digits_ratio = sum(c.isdigit() for c in normalized) / max(len(normalized), 1)
        if digits_ratio > 0.5:
            return True

        # Skip entities containing duration patterns ("X years", "X months", "X days")
        if re.match(r"^\d+\s+(?:year|month|day|week|hour|minute|second)s?$", normalized):
            return True

        # Skip date-like strings
        if re.match(r"^\d+[,\s]+\d{4}$", normalized):  # e.g. "16, 2025"
            return True
        if re.match(r"^(?:fiscal\s+)?(?:20|19)\d{2}[\-\s]?(?:20|19)?\d{0,2}$", normalized):
            return True  # "2024-2025", "2025"

        # --- PERSON-specific filters ---
        if label == "PERSON":
            # Must have at least 2 words for a full name
            words = text.split()
            if len(words) < 2:
                return True
            # Strictly enforce capitalization on all words (e.g. discard "a Bid")
            if not all(w[0].isupper() for w in words if w):
                return True
            # Skip if any word is a known non-name token
            non_name_words = {
                "limited", "private", "ltd", "pvt", "llp", "inc", "corp",
                "bank", "trust", "fund", "company", "group", "committee",
                "account", "form", "type", "category", "class", "series",
                "authorised", "authorized", "share", "capital", "paid",
                "act", "branch", "shares", "amount", "trusts", "slip",
                "schedule", "bid", "bids", "id", "dp", "village", "taluka",
                "allotted", "bidder", "bidders", "east", "mumbai", "pune", "bandra",
                "manager", "managers", "individual", "defaulter", "defaulters",
                "wilful", "circuit", "kilometers", "kilometer", "conditioning",
                "air", "running", "lead", "qib", "options", "option", "shoe",
                "green", "gram", "jyoti", "operational", "objects", "auditor",
                "auditors", "parent", "parents",
                "road", "lane", "photo", "voltaic", "pat", "margin", "mega",
                "volt", "amperes", "voltage", "ampere", "watt", "watts", "kw",
                "mw", "kva", "mva", "solar", "power", "grid", "plant", "plants",
                "energy", "electricity", "listing", "sebi", "bhavan", "registrar",
                "facility", "facilities", "paper", "papers", "news", "daily", "newspaper",
                "newspapers", "circular", "circulars", "prospectus", "memorandum",
                "application", "applications", "issue", "offer", "bidding", "anchor",
                "investor", "investors", "regulatory", "authority", "authorities",
                "government", "central", "state", "corporate", "commercial", "industrial",
                "financial", "performance", "indicator", "indicators", "metric", "metrics",
                "ratio", "ratios", "revenue", "revenues", "income", "profit", "loss",
                "tax", "taxes", "duties", "stamp", "registration", "depository",
                "depositories", "merchant", "banker", "bankers", "syndicate", "broker",
                "brokers", "widely", "circulated", "marathi"
            }
            if any(w.lower() in non_name_words for w in words):
                return True
            # Skip if text contains special chars typical of non-names
            if any(c in normalized for c in "@#$%^&*(){}[]|/\\<>~"):
                return True

        # --- ORG-specific filters ---
        if label == "ORG":
            # Skip single-word ORGs shorter than 4 chars
            if len(normalized.split()) == 1 and len(normalized) < 4:
                return True
            # Skip financial jargon falsely tagged as ORG
            org_noise = {
                "asba", "upi", "neft", "rtgs", "imps", "ecs",
                "pan", "tan", "din", "cin", "gstin", "gst",
                "neft", "nach", "mandate",
                "red herring prospectus", "offer document",
                "bid cum application form",
                "non-gaap measures", "non-gaap",
                "restated financial statements",
            }
            if normalized in org_noise:
                return True
            # Skip if entity text starts with common non-org starters
            noise_starters = (
                "the ", "a ", "an ", "our ", "their ", "its ",
                "section ", "chapter ", "part ", "schedule ",
                "pursuant ", "subject ", "including ",
            )
            if any(normalized.startswith(s) for s in noise_starters):
                return True

        return False

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 50_000, overlap: int = 500) -> List[Tuple[int, str]]:
        """Split text into overlapping chunks for NER processing."""
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            # Try to break at a sentence boundary
            if end < len(text):
                for sep in ["\n\n", "\n", ". ", "! ", "? "]:
                    idx = text.rfind(sep, start + chunk_size - overlap, end)
                    if idx > start:
                        end = idx + len(sep)
                        break
            chunks.append((start, text[start:end]))
            start = end - overlap if end < len(text) else end
        return chunks


# ---------------------------------------------------------------------------
# Address Detector (Heuristic)
# ---------------------------------------------------------------------------
class AddressDetector:
    """
    Detects Indian physical/mailing addresses using heuristic patterns.
    Looks for PIN codes and surrounding address-like text.
    """

    # Indian PIN code pattern with optional space
    PIN_PATTERN = re.compile(
        r"\b([1-9]\d{2})\s?(\d{3})\b"
    )

    # Address keywords
    ADDRESS_KEYWORDS = re.compile(
        r"\b(?:Flat|Floor|Plot|Building|Tower|Office|Suite|House|Village|"
        r"Taluka|District|Block|Sector|Road|Street|Lane|Nagar|Colony|"
        r"Marg|Complex|Centre|Center|Park|Estate|Apartment|"
        r"No\.|S\.No\.|Survey|Gat)\b",
        re.IGNORECASE,
    )

    # Indian state names
    STATES = re.compile(
        r"\b(?:Maharashtra|Karnataka|Tamil Nadu|Gujarat|Rajasthan|"
        r"Madhya Pradesh|Uttar Pradesh|West Bengal|Andhra Pradesh|"
        r"Telangana|Kerala|Bihar|Punjab|Haryana|Odisha|Jharkhand|"
        r"Chhattisgarh|Assam|Goa|Himachal Pradesh|Uttarakhand|"
        r"Jammu|Kashmir|Sikkim|Meghalaya|Mizoram|Manipur|Nagaland|"
        r"Tripura|Arunachal Pradesh)\b",
        re.IGNORECASE,
    )

    def detect(self, text: str) -> List[PIIEntity]:
        """Detect address blocks containing PIN codes and address markers."""
        entities = []
        seen_ranges = set()

        for match in self.PIN_PATTERN.finditer(text):
            pin = match.group(1) + match.group(2)
            # Check if this looks like a real Indian PIN (starts 1-9, 6 digits)
            if not self._is_valid_pin(pin):
                continue

            # Look backwards and forwards from PIN for address context
            search_start = max(0, match.start() - 200)
            search_end = min(len(text), match.end() + 50)
            context = text[search_start:search_end]

            # Check if context contains address keywords
            if self.ADDRESS_KEYWORDS.search(context) or self.STATES.search(context):
                # Find the address block boundaries
                addr_start, addr_end = self._find_address_bounds(
                    text, match.start(), match.end()
                )

                # Skip if overlapping with already detected address
                range_key = (addr_start, addr_end)
                if range_key in seen_ranges:
                    continue
                seen_ranges.add(range_key)

                addr_text = text[addr_start:addr_end].strip()
                if len(addr_text) > 15:  # Minimum address length
                    entities.append(PIIEntity(
                        start=addr_start,
                        end=addr_end,
                        entity_type="ADDRESS",
                        original_text=addr_text,
                        source="heuristic",
                        confidence=0.75,
                    ))

        return entities

    def _is_valid_pin(self, pin: str) -> bool:
        """Check if a 6-digit number is a plausible Indian PIN code."""
        if len(pin) != 6:
            return False
        # Indian PINs range from 110001 to 855117
        val = int(pin)
        return 100000 <= val <= 899999

    def _find_address_bounds(self, text: str, pin_start: int, pin_end: int) -> Tuple[int, int]:
        """Find the start and end of an address block around a PIN code."""
        # Look backwards for address start (comma, semicolon, newline, or section header)
        start = pin_start
        for i in range(pin_start - 1, max(0, pin_start - 300), -1):
            if text[i] in "\n":
                start = i + 1
                break
            if text[i] == ";" or text[i:i+2] == ": ":
                start = i + 1
                break
        else:
            start = max(0, pin_start - 200)

        # Look forwards past PIN for state name and "India"
        end = pin_end
        remainder = text[pin_end:pin_end + 100]
        # Include state and country if present
        state_match = self.STATES.search(remainder)
        if state_match:
            end = pin_end + state_match.end()
            # Check for ", India" after state
            india_match = re.search(r",?\s*India\b", text[end:end + 20])
            if india_match:
                end += india_match.end()

        return start, end


# ---------------------------------------------------------------------------
# Contact Person Name Detector (Heuristic)
# ---------------------------------------------------------------------------
class ContactPersonNameDetector:
    """
    Extracts person names following headers like 'Contact Person:' in contact lists.
    This helps capture names formatted with slashes or non-standard punctuation that spaCy misses.
    """

    HEADER_PATTERN = re.compile(
        r'\b(?:Contact Person|Compliance Officer|Company Secretary)\b\s*[:\-]\s*(.*?)(?=\b(?:Website|Telephone|E-mail|Email|Tel|Fax|Web|Contact|Compliance|Company|CS|CO)\b|\n|;|$)',
        re.IGNORECASE
    )
    NAME_PATTERN = re.compile(r'\b[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*){1,3}\b')

    def detect(self, text: str) -> List[PIIEntity]:
        entities = []
        non_name_words = {
            "limited", "private", "ltd", "pvt", "llp", "inc", "corp",
            "bank", "trust", "fund", "company", "group", "committee",
            "account", "form", "type", "category", "class", "series",
            "authorised", "authorized", "share", "capital", "paid",
            "act", "branch", "shares", "amount", "trusts", "slip",
            "schedule", "bid", "bids", "id", "dp", "village", "taluka",
            "allotted", "bidder", "bidders", "east", "mumbai", "pune", "bandra",
            "manager", "managers", "individual", "defaulter", "defaulters",
            "wilful", "circuit", "kilometers", "kilometer", "conditioning",
            "air", "running", "lead", "qib", "options", "option", "shoe",
            "green", "gram", "jyoti", "operational", "objects", "auditor",
            "auditors", "parent", "parents", "registrar",
            "road", "lane", "photo", "voltaic", "pat", "margin", "mega",
            "volt", "amperes", "voltage", "ampere", "watt", "watts", "kw",
            "mw", "kva", "mva", "solar", "power", "grid", "plant", "plants",
            "energy", "electricity", "listing", "sebi", "bhavan",
            "facility", "facilities", "paper", "papers", "news", "daily", "newspaper",
            "newspapers", "circular", "circulars", "prospectus", "memorandum",
            "application", "applications", "issue", "offer", "bidding", "anchor",
            "investor", "investors", "regulatory", "authority", "authorities",
            "government", "central", "state", "corporate", "commercial", "industrial",
            "financial", "performance", "indicator", "indicators", "metric", "metrics",
            "ratio", "ratios", "revenue", "revenues", "income", "profit", "loss",
            "tax", "taxes", "duties", "stamp", "registration", "depository",
            "depositories", "merchant", "banker", "bankers", "syndicate", "broker",
            "brokers", "widely", "circulated", "marathi"
        }
        for match in self.HEADER_PATTERN.finditer(text):
            content = match.group(1).strip()
            content_start = match.start(1)
            for name_match in self.NAME_PATTERN.finditer(content):
                name = name_match.group()
                words = name.lower().split()
                if any(w in non_name_words for w in words):
                    continue
                start = content_start + name_match.start()
                end = content_start + name_match.end()
                entities.append(PIIEntity(
                    start=start,
                    end=end,
                    entity_type="PERSON",
                    original_text=name,
                    source="heuristic",
                    confidence=0.90
                ))
        return entities


# ---------------------------------------------------------------------------
# Entity Resolution & Deduplication
# ---------------------------------------------------------------------------
def resolve_overlapping_entities(entities: List[PIIEntity]) -> List[PIIEntity]:
    """
    Resolve overlapping entity spans.
    Priority: higher confidence wins; on tie, more specific type wins.
    """
    if not entities:
        return []

    # Sort by start position, then by span length (longer first)
    entities.sort(key=lambda e: (e.start, -(e.end - e.start)))

    # Priority order for entity types (higher = preferred)
    type_priority = {
        "EMAIL": 10,
        "PHONE": 9,
        "SSN": 9,
        "CREDIT_CARD": 9,
        "IP_ADDRESS": 9,
        "CIN": 8,
        "DOB": 7,
        "PERSON": 6,
        "ORG": 5,
        "ADDRESS": 4,
        "GPE": 3,
        "LOC": 3,
        "DATE": 2,
        "URL": 6,
    }

    resolved = []
    for entity in entities:
        overlaps = False
        for existing in resolved:
            if entity.start < existing.end and entity.end > existing.start:
                # Overlap detected - keep higher priority
                e_priority = type_priority.get(entity.entity_type, 0)
                x_priority = type_priority.get(existing.entity_type, 0)
                if e_priority > x_priority or (
                    e_priority == x_priority and entity.confidence > existing.confidence
                ):
                    resolved.remove(existing)
                    resolved.append(entity)
                overlaps = True
                break
        if not overlaps:
            resolved.append(entity)

    return sorted(resolved, key=lambda e: e.start)


# ---------------------------------------------------------------------------
# DOCX Run-Level Replacement Engine
# ---------------------------------------------------------------------------
class DocxRedactor:
    """
    Handles PII replacement in .docx files at the XML run level,
    preserving formatting (bold, italic, font, color, etc.).
    """

    def __init__(self, anonymizer: ConsistentAnonymizer, stats: RedactionStats, ignored_types: Optional[Set[str]] = None):
        self.anonymizer = anonymizer
        self.stats = stats
        self.ignored_types = ignored_types or set()

    def redact_paragraph(self, paragraph, entities: List[PIIEntity]):
        """Replace PII entities in a paragraph while preserving run formatting."""
        # Filter out ignored types
        entities = [e for e in entities if e.entity_type not in self.ignored_types]

        if not entities or not paragraph.runs:
            return

        full_text = paragraph.text

        # Build character -> (run_idx, char_offset) mapping
        char_map = []
        for r_idx, run in enumerate(paragraph.runs):
            for c_offset in range(len(run.text)):
                char_map.append((r_idx, c_offset))

        if not char_map:
            return

        # Process entities in reverse order to preserve character offsets
        sorted_entities = sorted(entities, key=lambda e: e.start, reverse=True)

        for entity in sorted_entities:
            start = entity.start
            end = entity.end

            if start < 0 or end > len(char_map) or start >= end:
                continue

            # Get replacement text
            replacement = self.anonymizer.get_replacement(
                entity.original_text, entity.entity_type
            )
            entity.replacement = replacement
            self.stats.record(entity)

            start_run_idx, start_run_offset = char_map[start]
            end_run_idx, end_run_offset = char_map[end - 1]

            if start_run_idx == end_run_idx:
                # Single run replacement
                run = paragraph.runs[start_run_idx]
                run.text = (
                    run.text[:start_run_offset]
                    + replacement
                    + run.text[end_run_offset + 1:]
                )
            else:
                # Multi-run replacement
                first_run = paragraph.runs[start_run_idx]
                first_run.text = first_run.text[:start_run_offset] + replacement

                for mid_idx in range(start_run_idx + 1, end_run_idx):
                    paragraph.runs[mid_idx].text = ""

                last_run = paragraph.runs[end_run_idx]
                last_run.text = last_run.text[end_run_offset + 1:]

    def redact_text_simple(self, text: str, entities: List[PIIEntity]) -> str:
        """Simple string-level replacement for non-run contexts."""
        if not entities:
            return text

        sorted_entities = sorted(entities, key=lambda e: e.start, reverse=True)
        result = text

        for entity in sorted_entities:
            replacement = self.anonymizer.get_replacement(
                entity.original_text, entity.entity_type
            )
            entity.replacement = replacement
            self.stats.record(entity)
            result = result[:entity.start] + replacement + result[entity.end:]

        return result


# ---------------------------------------------------------------------------
# Main PII Detection Pipeline
# ---------------------------------------------------------------------------
class PIIDetectionPipeline:
    """
    Orchestrates multiple PII detectors and resolves conflicts.
    """

    def __init__(self, ner_model: Optional[str] = None):
        self.regex_detector = RegexDetector()
        self.ner_detector = NERDetector(model_name=ner_model)
        self.address_detector = AddressDetector()
        self.contact_detector = ContactPersonNameDetector()

    def detect_all(self, text: str) -> List[PIIEntity]:
        """Run all detectors and merge results."""
        all_entities = []

        # 1. Regex detection (highest priority for structured PII)
        regex_entities = self.regex_detector.detect(text)
        all_entities.extend(regex_entities)
        logger.debug(f"  Regex found {len(regex_entities)} entities")

        # 2. NER detection (for names, organizations, locations)
        ner_entities = self.ner_detector.detect(text)
        all_entities.extend(ner_entities)
        logger.debug(f"  NER found {len(ner_entities)} entities")

        # 3. Address detection (heuristic)
        addr_entities = self.address_detector.detect(text)
        all_entities.extend(addr_entities)
        logger.debug(f"  Address heuristic found {len(addr_entities)} entities")

        # 4. Contact person name detection (heuristic)
        contact_entities = self.contact_detector.detect(text)
        all_entities.extend(contact_entities)
        logger.debug(f"  Contact heuristic found {len(contact_entities)} entities")

        # Resolve overlaps
        resolved = resolve_overlapping_entities(all_entities)
        return resolved


# ---------------------------------------------------------------------------
# Main Redaction Engine
# ---------------------------------------------------------------------------
class PIIRedactionEngine:
    """
    End-to-end PII redaction engine for .docx documents.
    """

    def __init__(self, ner_model: Optional[str] = None, locale: str = "en_US",
                 salt: str = "pii_redaction_salt_2026"):
         self.pipeline = PIIDetectionPipeline(ner_model=ner_model)
         self.anonymizer = ConsistentAnonymizer(locale=locale, salt=salt)
         self.stats = RedactionStats()

    def redact_document(self, input_path: str, output_path: str, ignored_types: Optional[Set[str]] = None) -> RedactionStats:
        """
        Read a .docx document, detect and replace PII, save redacted version.
        """
        logger.info(f"Reading document: {input_path}")
        doc = Document(input_path)
        redactor = DocxRedactor(self.anonymizer, self.stats, ignored_types)

        # ---- Process Body Paragraphs ----
        logger.info("Processing body paragraphs...")
        para_count = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                entities = self.pipeline.detect_all(para.text)
                if entities:
                    redactor.redact_paragraph(para, entities)
            if (i + 1) % 100 == 0:
                logger.info(f"  Processed {i + 1}/{para_count} paragraphs")

        # ---- Process Tables ----
        logger.info(f"Processing {len(doc.tables)} tables...")
        for t_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            entities = self.pipeline.detect_all(para.text)
                            if entities:
                                redactor.redact_paragraph(para, entities)
            if (t_idx + 1) % 10 == 0:
                logger.info(f"  Processed {t_idx + 1}/{len(doc.tables)} tables")

        # ---- Process Headers & Footers ----
        logger.info("Processing headers and footers...")
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf is not None:
                    for para in hf.paragraphs:
                        if para.text.strip():
                            entities = self.pipeline.detect_all(para.text)
                            if entities:
                                redactor.redact_paragraph(para, entities)

        # ---- Save ----
        logger.info(f"Saving redacted document: {output_path}")
        doc.save(output_path)

        # ---- Report ----
        self._print_summary()
        self._save_mapping(output_path)

        return self.stats

    def _print_summary(self):
        """Print a summary of the redaction process."""
        logger.info("=" * 60)
        logger.info("REDACTION SUMMARY")
        logger.info("=" * 60)
        logger.info(f"Total PII entities redacted: {self.stats.total_entities}")
        logger.info("By PII Type:")
        for pii_type, count in sorted(self.stats.by_type.items()):
            logger.info(f"  {pii_type:20s}: {count}")
        logger.info("By Detection Source:")
        for source, count in sorted(self.stats.by_source.items()):
            logger.info(f"  {source:20s}: {count}")
        logger.info("=" * 60)

    def _save_mapping(self, output_path: str):
        """Save the PII replacement mapping to a JSON file."""
        mapping_path = Path(output_path).with_suffix(".mapping.json")
        mapping_data = {
            "total_entities": self.stats.total_entities,
            "by_type": self.stats.by_type,
            "by_source": self.stats.by_source,
            "replacements": self.stats.replacement_map,
        }
        with open(mapping_path, "w", encoding="utf-8") as f:
            json.dump(mapping_data, f, indent=2, ensure_ascii=False)
        logger.info(f"Replacement mapping saved: {mapping_path}")


# ---------------------------------------------------------------------------
# CLI Entry Point
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="PII Redaction Tool - Detect and replace PII in .docx documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python pii_redactor.py "Red Herring Prospectus.docx"
  python pii_redactor.py input.docx -o redacted_output.docx
  python pii_redactor.py input.docx --model en_core_web_trf --verbose
        """,
    )
    parser.add_argument(
        "input_file",
        help="Path to the input .docx file",
    )
    parser.add_argument(
        "-o", "--output",
        help="Path for the redacted output .docx file (default: <input>_redacted.docx)",
        default=None,
    )
    parser.add_argument(
        "--model",
        help="spaCy model to use for NER (default: en_core_web_sm)",
        default="en_core_web_sm",
    )
    parser.add_argument(
        "--locale",
        help="Faker locale for generating fake data (default: en_US)",
        default="en_US",
    )
    parser.add_argument(
        "--salt",
        help="Salt for deterministic fake data generation",
        default="pii_redaction_salt_2026",
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose/debug logging",
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger("pii_redactor").setLevel(logging.DEBUG)

    # Determine output path
    input_path = Path(args.input_file)
    if args.output:
        output_path = args.output
    else:
        output_path = str(input_path.with_name(f"{input_path.stem}_redacted{input_path.suffix}"))

    # Run redaction
    engine = PIIRedactionEngine(
        ner_model=args.model,
        locale=args.locale,
        salt=args.salt,
    )
    engine.redact_document(str(input_path), output_path)

    logger.info("Done! Redacted document saved.")


if __name__ == "__main__":
    main()
